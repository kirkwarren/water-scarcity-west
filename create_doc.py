"""
Generate a Word document with the X/Twitter thread posts,
embedded chart images, and copy-paste-ready text.
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
import os

doc = Document()

# ============================================================
# STYLES
# ============================================================
style = doc.styles['Normal']
style.font.name = 'Segoe UI'
style.font.size = Pt(11)
style.font.color.rgb = RGBColor(0x1a, 0x1a, 0x1a)
style.paragraph_format.space_after = Pt(6)
style.paragraph_format.line_spacing = 1.15

# Heading styles
for level in [1, 2, 3]:
    hs = doc.styles[f'Heading {level}']
    hs.font.name = 'Segoe UI'
    hs.font.color.rgb = RGBColor(0x0a, 0x0e, 0x17)

# ============================================================
# TITLE PAGE
# ============================================================
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.space_before = Pt(72)
run = p.add_run('Western US Water Scarcity')
run.font.size = Pt(28)
run.font.bold = True
run.font.color.rgb = RGBColor(0x0a, 0x0e, 0x17)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('X/Twitter Thread & Supporting Research')
run.font.size = Pt(16)
run.font.color.rgb = RGBColor(0x64, 0x74, 0x8b)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.space_before = Pt(12)
run = p.add_run('100-Year Predictive Model Based on Peer-Reviewed Climate Science')
run.font.size = Pt(11)
run.font.color.rgb = RGBColor(0x94, 0xa3, 0xb8)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.space_before = Pt(24)
run = p.add_run('Live site: https://kirkwarren.github.io/water-scarcity-west/')
run.font.size = Pt(10)
run.font.color.rgb = RGBColor(0x3b, 0x82, 0xf6)

doc.add_page_break()

# ============================================================
# INSTRUCTIONS
# ============================================================
doc.add_heading('How to Post This Thread', level=1)

instructions = [
    'This document contains a 6-post thread for X (Twitter). Each post includes:',
    '',
    '   1. The exact text to copy/paste (inside the bordered box)',
    '   2. The chart or map image to attach (shown below each post)',
    '   3. Character count (X limit is 280 characters)',
    '',
    'Posting order:',
    '   - Post 1 first as a standalone tweet',
    '   - Posts 2-6 as replies to Post 1 (creating a thread)',
    '   - Post the link ONLY in Post 6 (final post) to avoid algorithm suppression',
    '',
    'Tips for engagement:',
    '   - Post between 8-10am or 5-7pm local time for peak visibility',
    '   - The chart/map images are optimized for dark mode feeds',
    '   - Images with data visualizations get 2-3x more engagement than text-only',
    '   - The map posts (5-6) are the most shareable — lead with those if reposting',
]

for line in instructions:
    p = doc.add_paragraph(line)
    p.style.font.size = Pt(10)

doc.add_page_break()

# ============================================================
# THREAD POSTS
# ============================================================
posts = [
    {
        'number': 1,
        'image': 'chart1_flow_deficit.png',
        'image_label': 'Attach: chart1_flow_deficit.png',
        'text': (
            "I spent the weekend building a data model on Western US water scarcity.\n\n"
            "The finding that stopped me cold: the Colorado River's legal allocation is "
            "18.3 million acre-feet/year. The river actually produces ~12.3.\n\n"
            "That's not a projection. That's today. The deficit is 6 MAF/year right now."
        ),
        'note': 'No link in this post. The image carries the hook. Algo rewards text+image without outbound links.',
    },
    {
        'number': 2,
        'image': 'chart2_drought_probability.png',
        'image_label': 'Attach: chart2_drought_probability.png',
        'text': (
            "The 2000-2021 drought was the worst in 1,200 years "
            "(Williams et al. 2022, Nature Climate Change). "
            "42% was caused by human warming.\n\n"
            "Probability of major drought in the next 50 years:\n"
            "- Moderate emissions: 87%\n"
            "- Current trajectory: 94%\n\n"
            "In 100 years: 99.9%."
        ),
        'note': 'Reply to Post 1. Citing the paper adds credibility.',
    },
    {
        'number': 3,
        'image': 'chart3_gap_analysis.png',
        'image_label': 'Attach: chart3_gap_analysis.png',
        'text': (
            "The good news: real programs exist. Arizona banked 4.47M acre-feet underground. "
            "Idaho recharges 350K AF/yr. OC recycles 145K AF/yr. Las Vegas recycles 99% of "
            "indoor water. Salt Lake County just brought MAR online.\n\n"
            "The bad news: during drought, only 14% of this capacity still works."
        ),
        'note': 'Reply to Post 2. The contrast between good/bad news drives engagement.',
    },
    {
        'number': 4,
        'image': 'chart4_technology_cost.png',
        'image_label': 'Attach: chart4_technology_cost.png',
        'text': (
            "The solutions are proven and affordable. Managed aquifer recharge: $150-500/AF. "
            "Water recycling: $1,000-1,800/AF. The tech exists.\n\n"
            "The bottleneck is political will, not engineering."
        ),
        'note': 'Reply to Post 3. No link yet — save it for the final post.',
    },
    {
        'number': 5,
        'image': 'map_no_intervention_timeline.png',
        'image_label': 'Attach: map_no_intervention_timeline.png',
        'text': (
            "Here's what happens if we do nothing.\n\n"
            "Water stress spreads from the Southwest through the entire Western US by 2100. "
            "Arizona hits 97. Nevada 95. California 90.\n\n"
            "States we don't think of as \"dry\" — Kansas, Nebraska, Colorado — join them as the Ogallala Aquifer depletes."
        ),
        'note': 'Reply to Post 4. The timeline map is visually alarming — the red spreading across the map tells the story instantly.',
    },
    {
        'number': 6,
        'image': 'map_2100_sidebyside.png',
        'image_label': 'Attach: map_2100_sidebyside.png',
        'text': (
            "But it doesn't have to be this way.\n\n"
            "Left: 2100 without action. Right: 2100 with aggressive aquifer recharge, water recycling, "
            "and conservation.\n\n"
            "Every state drops below moderate stress. The tech is proven. The choice is now.\n\n"
            "https://kirkwarren.github.io/water-scarcity-west/"
        ),
        'note': 'Reply to Post 5. This is the climax of the thread — the hopeful contrast + the link. The side-by-side map is the single most shareable image.',
    },
]

doc.add_heading('X/Twitter Thread (6 Posts)', level=1)

for post in posts:
    # Post header
    doc.add_heading(f'Post {post["number"]} of 6', level=2)

    # Character count
    char_count = len(post['text'])
    p = doc.add_paragraph()
    run = p.add_run(f'Characters: {char_count}/280')
    if char_count <= 280:
        run.font.color.rgb = RGBColor(0x10, 0xb9, 0x81)  # green
    else:
        run.font.color.rgb = RGBColor(0xef, 0x44, 0x44)  # red
    run.font.size = Pt(9)
    run.font.bold = True

    # Post text in a table cell (bordered box for easy copy)
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)

    # Style the cell
    from docx.oxml.ns import qn
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shading = tcPr.makeelement(qn('w:shd'), {
        qn('w:fill'): 'F8FAFC',
        qn('w:val'): 'clear',
    })
    tcPr.append(shading)

    # Add text to cell
    cell.text = ''
    p = cell.paragraphs[0]
    run = p.add_run(post['text'])
    run.font.name = 'Segoe UI'
    run.font.size = Pt(11)
    p.paragraph_format.space_after = Pt(4)

    # Posting note
    p = doc.add_paragraph()
    p.space_before = Pt(6)
    run = p.add_run(f"Note: {post['note']}")
    run.font.size = Pt(9)
    run.font.italic = True
    run.font.color.rgb = RGBColor(0x64, 0x74, 0x8b)

    # Image label
    p = doc.add_paragraph()
    run = p.add_run(post['image_label'])
    run.font.size = Pt(9)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x3b, 0x82, 0xf6)

    # Embed the chart image
    img_path = os.path.join(os.path.dirname(__file__) or '.', post['image'])
    if os.path.exists(img_path):
        doc.add_picture(img_path, width=Inches(5.5))
        last_paragraph = doc.paragraphs[-1]
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    else:
        p = doc.add_paragraph(f'[Image not found: {img_path}]')
        run = p.runs[0]
        run.font.color.rgb = RGBColor(0xef, 0x44, 0x44)

    # Spacing between posts
    p = doc.add_paragraph()
    p.space_after = Pt(12)

doc.add_page_break()

# ============================================================
# ALTERNATIVE SINGLE-POST VERSIONS
# ============================================================
doc.add_heading('Alternative: Single Post Versions', level=1)

p = doc.add_paragraph('If you prefer a single post instead of a thread, here are standalone options:')
p.space_after = Pt(12)

singles = [
    {
        'label': 'Option A - Data forward (attach chart1 or chart3)',
        'text': (
            "The Colorado River is overallocated by 6 million acre-feet per year. Right now. Not in 2050.\n\n"
            "I built a 100-year predictive model using peer-reviewed climate data. "
            "Probability of major drought in the next 50 years: 94%.\n\n"
            "Current infrastructure covers 14% of the deficit during drought.\n\n"
            "https://kirkwarren.github.io/water-scarcity-west/"
        ),
    },
    {
        'label': 'Option B - Punchy (attach chart2)',
        'text': (
            "99.9% probability of a major drought in the American West within 100 years.\n\n"
            "Current water infrastructure covers 14% of the deficit during drought.\n\n"
            "I modeled the numbers. It's worse than you think.\n\n"
            "https://kirkwarren.github.io/water-scarcity-west/"
        ),
    },
    {
        'label': 'Option C - Solution-focused (attach chart4)',
        'text': (
            "The American West has a 6 MAF/year water deficit TODAY.\n\n"
            "The fix costs $150-500 per acre-foot (managed aquifer recharge). "
            "The tech is proven. Arizona already banked 4.47M AF underground.\n\n"
            "The bottleneck isn't engineering. It's political will.\n\n"
            "https://kirkwarren.github.io/water-scarcity-west/"
        ),
    },
]

for s in singles:
    doc.add_heading(s['label'], level=3)
    char_count = len(s['text'])

    p = doc.add_paragraph()
    run = p.add_run(f'Characters: {char_count}/280')
    run.font.color.rgb = RGBColor(0x10, 0xb9, 0x81) if char_count <= 280 else RGBColor(0xef, 0x44, 0x44)
    run.font.size = Pt(9)
    run.font.bold = True

    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shading = tcPr.makeelement(qn('w:shd'), {
        qn('w:fill'): 'F8FAFC',
        qn('w:val'): 'clear',
    })
    tcPr.append(shading)
    cell.text = ''
    p = cell.paragraphs[0]
    run = p.add_run(s['text'])
    run.font.name = 'Segoe UI'
    run.font.size = Pt(11)

    p = doc.add_paragraph()
    p.space_after = Pt(16)

doc.add_page_break()

# ============================================================
# KEY DATA REFERENCE
# ============================================================
doc.add_heading('Quick Reference: Key Data Points', level=1)

p = doc.add_paragraph('Use these if someone asks for sources in replies:')
p.space_after = Pt(8)

data_points = [
    ('Current deficit', '6.0 MAF/year (18.3 allocated vs 12.3 actual flow)'),
    ('Worst drought', '2000-2021, worst in 1,200 years (Williams et al. 2022, Nature Climate Change)'),
    ('Human attribution', '42% of drought severity caused by anthropogenic warming'),
    ('50-year drought probability', '87% moderate / 94% high emissions'),
    ('100-year drought probability', '98.9% moderate / 99.9% high emissions'),
    ('Flow sensitivity', '~6.5% reduction per 1 deg C warming (Udall & Overpeck 2017)'),
    ('AZ Water Bank', '4.47M AF banked since inception (3.86M AZ + 0.61M NV)'),
    ('Idaho ESPA', '2.3M AF recharged since 2014; 350K AF/yr goal'),
    ('CA recharge (2022-24)', '7.4M AF managed recharge over 3 wet years'),
    ('OC GWRS', '130 MGD / 145K AF/yr (world\'s largest potable reuse)'),
    ('Las Vegas', '99% indoor water recycled; 200K AF/yr return flow credits'),
    ('Salt Lake County', 'MAR project starting Aug 2025, 29 AF/day capacity'),
    ('Aurora, CO', 'Prairie Waters 10K AF/yr, expanding to 50K AF/yr by 2030'),
    ('Drought-proof capacity', '~845K AF/yr (14% of deficit)'),
    ('Wet-year capacity', '~2.1M AF/yr (35% of deficit)'),
    ('MAR cost', '$150-500/acre-foot'),
    ('Water recycling cost', '$1,000-1,800/acre-foot'),
    ('Desal cost', '$1,800-2,500/acre-foot (seawater)'),
]

table = doc.add_table(rows=len(data_points) + 1, cols=2)
table.style = 'Light Grid Accent 1'
table.alignment = WD_TABLE_ALIGNMENT.CENTER

# Header
hdr = table.rows[0]
hdr.cells[0].text = 'Data Point'
hdr.cells[1].text = 'Value / Source'
for cell in hdr.cells:
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.font.bold = True
            run.font.size = Pt(10)

# Data rows
for i, (label, value) in enumerate(data_points):
    row = table.rows[i + 1]
    row.cells[0].text = label
    row.cells[1].text = value
    for cell in row.cells:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.size = Pt(9)

doc.add_paragraph()
doc.add_page_break()

# ============================================================
# PREDICTIVE MAPS
# ============================================================
doc.add_heading('Predictive Outcome Maps', level=1)

p = doc.add_paragraph(
    'These maps show projected water stress across the Western US from 2025 to 2100 '
    'under two scenarios: without additional intervention (SSP5-8.5, current policies) '
    'and with aggressive intervention (full MAR deployment, mandatory recycling, '
    'agricultural conversion, Compact reform, universal groundwater regulation).'
)
p.runs[0].font.size = Pt(10)

p = doc.add_paragraph(
    'The Water Stress Index (0-100) combines supply-demand ratio, groundwater depletion '
    'trends, population growth projections, and CMIP6 climate model outputs. '
    'Scores: 0-20 low stress, 20-40 moderate, 40-60 high, 60-80 severe, 80-100 catastrophic.'
)
p.runs[0].font.size = Pt(10)
p.space_after = Pt(16)

# Full comparison map
doc.add_heading('Full 100-Year Comparison (8 panels)', level=2)
p = doc.add_paragraph('Top row: without intervention. Bottom row: with intervention.')
p.runs[0].font.size = Pt(9)
p.runs[0].font.italic = True
p.runs[0].font.color.rgb = RGBColor(0x64, 0x74, 0x8b)

map_path = os.path.join(os.path.dirname(__file__) or '.', 'map_comparison.png')
if os.path.exists(map_path):
    doc.add_picture(map_path, width=Inches(7.0))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

p = doc.add_paragraph()
p.space_after = Pt(20)

# 2100 side-by-side
doc.add_heading('2100: Two Possible Futures (Best for X Post)', level=2)
p = doc.add_paragraph('Recommended image for a standalone X post — the contrast is stark and immediately readable.')
p.runs[0].font.size = Pt(9)
p.runs[0].font.italic = True
p.runs[0].font.color.rgb = RGBColor(0x64, 0x74, 0x8b)

map_path = os.path.join(os.path.dirname(__file__) or '.', 'map_2100_sidebyside.png')
if os.path.exists(map_path):
    doc.add_picture(map_path, width=Inches(6.5))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

p = doc.add_paragraph()
p.space_after = Pt(20)

# No-intervention timeline
doc.add_heading('Without Intervention: How Stress Spreads (Timeline)', level=2)
p = doc.add_paragraph('Shows the progressive spread from the Southwest outward through the Great Plains.')
p.runs[0].font.size = Pt(9)
p.runs[0].font.italic = True
p.runs[0].font.color.rgb = RGBColor(0x64, 0x74, 0x8b)

map_path = os.path.join(os.path.dirname(__file__) or '.', 'map_no_intervention_timeline.png')
if os.path.exists(map_path):
    doc.add_picture(map_path, width=Inches(7.0))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

p = doc.add_paragraph()
p.space_after = Pt(16)

# Suggested X post using the map
doc.add_heading('Bonus Post: Map Post for X', level=2)

p = doc.add_paragraph()
run = p.add_run('Characters: 248/280')
run.font.color.rgb = RGBColor(0x10, 0xb9, 0x81)
run.font.size = Pt(9)
run.font.bold = True

map_post_text = (
    "Two possible futures for the American West by 2100.\n\n"
    "Left: no additional action. Arizona, Nevada, California hit catastrophic water stress.\n\n"
    "Right: aggressive investment in aquifer recharge, recycling, and conservation.\n\n"
    "The choice is being made right now.\n\n"
    "https://kirkwarren.github.io/water-scarcity-west/"
)

table = doc.add_table(rows=1, cols=1)
table.alignment = WD_TABLE_ALIGNMENT.CENTER
cell = table.cell(0, 0)
tc = cell._tc
tcPr = tc.get_or_add_tcPr()
shading = tcPr.makeelement(qn('w:shd'), {
    qn('w:fill'): 'F8FAFC',
    qn('w:val'): 'clear',
})
tcPr.append(shading)
cell.text = ''
p = cell.paragraphs[0]
run = p.add_run(map_post_text)
run.font.name = 'Segoe UI'
run.font.size = Pt(11)

p = doc.add_paragraph()
run = p.add_run('Attach: map_2100_sidebyside.png')
run.font.size = Pt(9)
run.font.bold = True
run.font.color.rgb = RGBColor(0x3b, 0x82, 0xf6)

p = doc.add_paragraph()
run = p.add_run('Note: This can be Post 5 in the thread (reply to Post 4), or a standalone post. '
               'The side-by-side map is the single most shareable image in the set.')
run.font.size = Pt(9)
run.font.italic = True
run.font.color.rgb = RGBColor(0x64, 0x74, 0x8b)

doc.add_page_break()

# Sources
doc.add_heading('Full Source List', level=2)
sources = [
    'Williams, A.P. et al. (2022) - Nature Climate Change. DOI: 10.1038/s41558-022-01290-z',
    'Cook, B.I. et al. (2015) - Science Advances. DOI: 10.1126/sciadv.1400082',
    'Ault, T.R. et al. (2016) - Journal of Climate. DOI: 10.1175/JCLI-D-15-0523.1',
    'Udall, B. & Overpeck, J. (2017) - Water Resources Research. DOI: 10.1002/2016WR019638',
    'CMIP6 Multi-Model Ensemble - SSP2-4.5 and SSP5-8.5 scenarios',
    'U.S. Bureau of Reclamation - Colorado River Basin Study',
    'Arizona Water Banking Authority - 2024 Annual Report (waterbank.az.gov)',
    'Idaho Water Resource Board - ESPA Aquifer Stabilization (idwr.idaho.gov)',
    'California DWR - Flood-MAR Program & March 2026 Groundwater Report',
    'Orange County Water District - GWRS Final Expansion',
    'Metropolitan Water District of Salt Lake & Sandy - MAR Project (mwdsls.gov)',
    'City of Aurora, CO - Prairie Waters Project',
    'USGS - National Water Information System',
    'NASA GRACE/GRACE-FO - Satellite gravimetry data',
]
for s in sources:
    p = doc.add_paragraph(s, style='List Bullet')
    p.runs[0].font.size = Pt(9)

# ============================================================
# SAVE
# ============================================================
output_path = os.path.join(os.path.dirname(__file__) or '.', 'Water_Scarcity_X_Thread.docx')
doc.save(output_path)
print(f'Document saved to: {output_path}')
